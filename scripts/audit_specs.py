# -*- coding: utf-8 -*-
"""Audit procurement spec documents for common dispute risks.

Usage:
  python scripts/audit_specs.py "C:\\path\\to\\folder"
  python scripts/audit_specs.py "C:\\path\\to\\file.docx"
"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path

try:
    import docx
except ImportError as exc:  # pragma: no cover
    raise SystemExit("python-docx is required. Install it before running this script.") from exc


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_RULES = ROOT / "config" / "spec-audit-rules.json"


@dataclass
class Finding:
    severity: str
    code: str
    message: str


def read_docx(path: Path) -> tuple[str, list[str]]:
    document = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    headings = [
        p.text.strip()
        for p in document.paragraphs
        if p.text.strip() and p.style and "Heading" in (p.style.name or "")
    ]

    table_text: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    table_text.append(value)

    return "\n".join(paragraphs + table_text), headings


def iter_targets(path: Path) -> list[Path]:
    if path.is_file():
        return [path] if path.suffix.lower() == ".docx" else []

    include_tokens = ("規格說明書", "規格書", "RFP", "需求規範", "需求說明", "招標文件-資格與規格")
    exclude_tokens = ("審查報告", "審查意見", "審查表", "請購申請單", "補充報告", "補充事項", "彙總報告")

    targets: list[Path] = []
    for item in path.rglob("*"):
        if item.is_file() and item.suffix.lower() == ".docx":
            name = item.name
            if any(token in name for token in include_tokens) and not any(token in name for token in exclude_tokens):
                targets.append(item)
    return sorted(targets, key=lambda p: str(p))


def contains_any(text: str, keywords: list[str]) -> bool:
    return any(keyword.lower() in text.lower() for keyword in keywords)


def audit_file(path: Path, base: Path, rules: dict) -> dict:
    findings: list[Finding] = []

    try:
        text, headings = read_docx(path)
    except Exception as exc:
        return {
            "path": str(path),
            "relative_path": str(path.relative_to(base)) if path.is_relative_to(base) else str(path),
            "score": 0,
            "headings": 0,
            "findings": [
                {
                    "severity": "HIGH",
                    "code": "READ_FAILED",
                    "message": f"無法讀取 docx：{exc}",
                }
            ],
        }

    if not re.match(rules["filename_pattern"], path.name):
        findings.append(
            Finding(
                "MEDIUM",
                "FILENAME",
                "檔名不符合 `{案號}_{文件類別}_{版次}_YYYYMMDD.docx`。",
            )
        )

    if not headings:
        findings.append(Finding("MEDIUM", "NO_HEADINGS", "未使用 Word Heading Styles，無法穩定產生目錄與自動審查。"))

    for section in rules["required_sections"]:
        if not contains_any(text, section["keywords"]):
            findings.append(Finding("HIGH", f"MISSING_{section['id'].upper()}", f"缺少必備章節或關鍵內容：{section['label']}。"))

    for conditional in rules["conditional_sections"]:
        triggered = contains_any(text, conditional["trigger_keywords"])
        if triggered:
            missing = [kw for kw in conditional["required_keywords"] if kw.lower() not in text.lower()]
            if missing:
                findings.append(
                    Finding(
                        "HIGH",
                        f"INCOMPLETE_{conditional['id'].upper()}",
                        f"{conditional['label']}內容不足，缺少：{', '.join(missing)}。",
                    )
                )

    ambiguous_hits = [term for term in rules["ambiguous_terms"] if term in text]
    if ambiguous_hits:
        findings.append(
            Finding(
                "MEDIUM",
                "AMBIGUOUS_TERMS",
                "發現容易造成爭議的模糊用語：" + ", ".join(sorted(set(ambiguous_hits))) + "。",
            )
        )

    high_count = sum(1 for item in findings if item.severity == "HIGH")
    medium_count = sum(1 for item in findings if item.severity == "MEDIUM")
    score = max(0, 100 - high_count * 12 - medium_count * 5)

    return {
        "path": str(path),
        "relative_path": str(path.relative_to(base)) if path.is_relative_to(base) else str(path),
        "score": score,
        "headings": len(headings),
        "findings": [item.__dict__ for item in findings],
    }


def print_markdown(results: list[dict]) -> None:
    print("# 規格書稽核結果")
    print()
    print("| 分數 | 高風險 | 中風險 | Heading | 文件 |")
    print("|------|--------|--------|---------|------|")
    for result in sorted(results, key=lambda item: (item["score"], item["relative_path"])):
        high = sum(1 for item in result["findings"] if item["severity"] == "HIGH")
        medium = sum(1 for item in result["findings"] if item["severity"] == "MEDIUM")
        print(f"| {result['score']} | {high} | {medium} | {result['headings']} | {result['relative_path']} |")

    print()
    print("## 詳細缺失")
    for result in sorted(results, key=lambda item: (item["score"], item["relative_path"])):
        if not result["findings"]:
            continue
        print()
        print(f"### {result['relative_path']}")
        for finding in result["findings"]:
            print(f"- [{finding['severity']}] {finding['code']}: {finding['message']}")


def main() -> int:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="docx file or folder to audit")
    parser.add_argument("--rules", default=str(DEFAULT_RULES), help="rules JSON path")
    parser.add_argument("--json", action="store_true", help="output JSON instead of Markdown")
    args = parser.parse_args()

    target = Path(args.path)
    rules_path = Path(args.rules)

    with rules_path.open("r", encoding="utf-8") as handle:
        rules = json.load(handle)

    targets = iter_targets(target)
    base = target if target.is_dir() else target.parent
    results = [audit_file(item, base, rules) for item in targets]

    if args.json:
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        print_markdown(results)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
