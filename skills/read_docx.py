# -*- coding: utf-8 -*-
"""Read a .docx procurement spec and extract structured content for review.

Usage:
    python skills/read_docx.py input.docx
    python skills/read_docx.py input.docx --json
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.exit("python-docx is required. Run: pip install python-docx")


def extract_structure(path: Path) -> dict:
    """Extract headings, paragraphs, and tables from a .docx file."""
    doc = Document(str(path))

    result = {
        "file": path.name,
        "headings": [],
        "paragraphs": [],
        "tables": [],
        "full_text": "",
    }

    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"
        full_text_parts.append(text)

        if "Heading" in style_name:
            level_match = re.search(r"(\d+)", style_name)
            level = int(level_match.group(1)) if level_match else 1
            result["headings"].append({"level": level, "text": text})
        else:
            result["paragraphs"].append({"style": style_name, "text": text})

    for i, table in enumerate(doc.tables):
        table_data = {"index": i, "rows": []}
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data["rows"].append(row_data)
        result["tables"].append(table_data)

    result["full_text"] = "\n".join(full_text_parts)
    return result


def match_chapters(structure: dict, schema: dict) -> dict:
    """Match extracted content against schema chapters."""
    chapters_found = {}
    full_text = structure["full_text"]
    headings_text = [h["text"] for h in structure["headings"]]

    for chapter in schema.get("chapters", []):
        ch_id = chapter["id"]
        ch_title = chapter["title"]
        ch_num = chapter["number"]

        # Check if chapter exists by heading or keyword
        found = False
        for h in headings_text:
            if ch_title in h or f"{ch_num}." in h:
                found = True
                break

        # Fallback: check full text for keywords
        if not found:
            keywords = [ch_title]
            if found := any(kw in full_text for kw in keywords):
                pass

        chapters_found[ch_id] = {
            "title": ch_title,
            "found": found,
            "required": chapter.get("required", False),
        }

    return chapters_found


def scan_ambiguous_terms(full_text: str, terms: list[str]) -> list[dict]:
    """Find ambiguous terms in the text."""
    findings = []
    lines = full_text.split("\n")
    for i, line in enumerate(lines, 1):
        for term in terms:
            if term in line:
                findings.append({
                    "line": i,
                    "term": term,
                    "context": line[:100],
                })
    return findings


def review(path: Path, schema_path: Path | None = None) -> dict:
    """Full review of a .docx file against schema."""
    structure = extract_structure(path)

    schema = {}
    if schema_path and schema_path.exists():
        with open(schema_path, encoding="utf-8") as f:
            schema = json.load(f)

    result = {
        "file": path.name,
        "heading_count": len(structure["headings"]),
        "table_count": len(structure["tables"]),
        "paragraph_count": len(structure["paragraphs"]),
    }

    if schema:
        result["chapters"] = match_chapters(structure, schema)
        result["ambiguous_terms"] = scan_ambiguous_terms(
            structure["full_text"],
            schema.get("ambiguous_terms", []),
        )

        # Checklist evaluation
        checklist_results = []
        for category in schema.get("checklist", {}).get("categories", []):
            cat_result = {"id": category["id"], "title": category["title"], "items": []}
            for item in category.get("items", []):
                # Simple heuristic: check if related keywords exist in text
                passed = any(
                    keyword in structure["full_text"]
                    for keyword in item["text"].split("（")[0].split("/")
                )
                cat_result["items"].append({
                    "id": item["id"],
                    "text": item["text"],
                    "blocking": item.get("blocking", False),
                    "passed": passed,
                })
            checklist_results.append(cat_result)
        result["checklist"] = checklist_results

    return result


def main():
    parser = argparse.ArgumentParser(description="Read and review .docx spec")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    parser.add_argument("--schema", default=None, help="Schema JSON path")
    args = parser.parse_args()

    path = Path(args.input)
    if not path.exists():
        sys.exit(f"File not found: {path}")

    schema_path = Path(args.schema) if args.schema else Path(__file__).parent / "spec-schema.json"

    result = review(path, schema_path)

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(f"檔案：{result['file']}")
        print(f"Heading 數：{result['heading_count']}")
        print(f"表格數：{result['table_count']}")
        print(f"段落數：{result['paragraph_count']}")

        if "chapters" in result:
            print("\n章節檢查：")
            for ch_id, info in result["chapters"].items():
                status = "✅" if info["found"] else ("❌" if info["required"] else "⚠️")
                print(f"  {status} {info['title']}")

        if "ambiguous_terms" in result and result["ambiguous_terms"]:
            print(f"\n模糊用語：找到 {len(result['ambiguous_terms'])} 處")
            for finding in result["ambiguous_terms"][:10]:
                print(f"  Line {finding['line']}: 「{finding['term']}」 → {finding['context']}")

        if "checklist" in result:
            print("\nChecklist：")
            total = 0
            passed = 0
            for cat in result["checklist"]:
                cat_pass = sum(1 for item in cat["items"] if item["passed"])
                cat_total = len(cat["items"])
                total += cat_total
                passed += cat_pass
                print(f"  {cat['id']}. {cat['title']}: {cat_pass}/{cat_total}")
            print(f"\n總分：{passed}/{total} ({passed/total*100:.0f}%)" if total else "")


if __name__ == "__main__":
    main()
