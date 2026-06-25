# -*- coding: utf-8 -*-
"""Generate a .docx procurement spec from structured data.

Usage:
    python skills/generate_docx.py input.json output.docx
    python skills/generate_docx.py --help

Input JSON format matches spec-schema.json structure with filled values.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx is required. Run: pip install python-docx")


def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_table(doc, columns: list[str], rows: list[list[str]], style: str = "Table Grid"):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(columns), style=style)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, col in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = col
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value) if value else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def build_cover_page(doc, data: dict):
    """Add cover page with case info."""
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data.get("case_name", "採購規格說明書"))
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()

    info_items = [
        ("案號", data.get("case_number", "")),
        ("文件類別", data.get("doc_type", "規格說明書")),
        ("文件版本", data.get("doc_version", "v1.0")),
        ("文件日期", data.get("doc_date", "")),
        ("文件狀態", data.get("doc_status", "Draft")),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(14)

    doc.add_page_break()


def add_toc(doc):
    """Add a Table of Contents field."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目　錄")
    run.bold = True
    run.font.size = Pt(16)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = run2._element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._element.append(fld_char_end)

    doc.add_paragraph("（請在 Word 中按右鍵 → 更新功能變數 以產生目錄）")
    doc.add_page_break()


def build_chapter(doc, chapter: dict, data: dict):
    """Build a single chapter from schema + data."""
    ch_num = chapter["number"]
    ch_title = chapter["title"]
    ch_data = data.get(chapter["id"], {})

    doc.add_heading(f"{ch_num}. {ch_title}", level=1)

    # Fields
    for field in chapter.get("fields", []):
        value = ch_data.get(field["id"], field.get("default", ""))
        if value:
            p = doc.add_paragraph()
            run = p.add_run(f"{field['label']}：")
            run.bold = True
            p.add_run(str(value))

    # Tables at chapter level
    for tbl in chapter.get("tables", []):
        if tbl.get("title"):
            doc.add_heading(tbl["title"], level=3)
        rows = ch_data.get(tbl["id"], tbl.get("fixed_rows", tbl.get("default_rows", [])))
        if rows:
            add_table(doc, tbl["columns"], rows)
        if tbl.get("note"):
            p = doc.add_paragraph()
            run = p.add_run(f"※ {tbl['note']}")
            run.italic = True
            run.font.size = Pt(9)

    # Sections
    for section in chapter.get("sections", []):
        sec_data = ch_data.get(section["id"], {})

        doc.add_heading(f"{ch_num}.{section.get('number', '')} {section['title']}".strip(), level=2)

        if section.get("note"):
            p = doc.add_paragraph()
            run = p.add_run(f"※ {section['note']}")
            run.italic = True
            run.font.size = Pt(9)

        # Section fields
        for field in section.get("fields", []):
            value = sec_data.get(field["id"], field.get("default", ""))
            if value:
                p = doc.add_paragraph()
                run = p.add_run(f"{field['label']}：")
                run.bold = True
                p.add_run(str(value))

        # Section tables
        for tbl in section.get("tables", []):
            if tbl.get("title"):
                doc.add_heading(tbl["title"], level=3)
            rows = sec_data.get(tbl["id"], tbl.get("fixed_rows", tbl.get("default_rows", [])))
            if rows:
                add_table(doc, tbl["columns"], rows)
            if tbl.get("hint"):
                p = doc.add_paragraph()
                run = p.add_run(f"提示：{tbl['hint']}")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Subsections
        for sub in section.get("subsections", []):
            sub_data = sec_data.get(sub["id"], {})
            doc.add_heading(sub["title"], level=3)

            for tbl in sub.get("tables", []):
                rows = sub_data.get(tbl["id"], tbl.get("fixed_rows", tbl.get("default_rows", [])))
                if rows:
                    add_table(doc, tbl["columns"], rows)

    doc.add_page_break()


def generate(schema: dict, data: dict, output_path: Path):
    """Generate the .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "微軟正黑體"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")

    # Set heading fonts
    for i in range(1, 4):
        h_style = doc.styles[f"Heading {i}"]
        h_style.font.name = "微軟正黑體"
        h_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微軟正黑體")

    # Cover page
    cover_data = data.get("ch0", {})
    build_cover_page(doc, cover_data)

    # TOC
    add_toc(doc)

    # Chapters
    for chapter in schema["chapters"]:
        # Skip conditional chapters if not applicable
        if chapter.get("conditional") and not data.get(chapter["id"]):
            continue
        build_chapter(doc, chapter, data)

    doc.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Generate .docx procurement spec")
    parser.add_argument("input", help="Input JSON file with spec data")
    parser.add_argument("output", help="Output .docx file path")
    parser.add_argument("--schema", default=None, help="Schema JSON (default: skills/spec-schema.json)")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    schema_path = Path(args.schema) if args.schema else Path(__file__).parent / "spec-schema.json"
    if not schema_path.exists():
        sys.exit(f"Schema not found: {schema_path}")

    with open(schema_path, encoding="utf-8") as f:
        schema = json.load(f)

    with open(input_path, encoding="utf-8") as f:
        data = json.load(f)

    result = generate(schema, data, output_path)
    print(f"Generated: {result}")


if __name__ == "__main__":
    main()
